import csv
import io
import json
import logging
from pathlib import Path

from fastapi import UploadFile

logger = logging.getLogger(__name__)


class UnsupportedDocumentError(ValueError):
    pass


class DocumentParserService:
    supported_suffixes = {".txt", ".md", ".csv", ".json", ".pdf", ".docx", ".xlsx"}

    async def extract_text(self, file: UploadFile) -> str:
        suffix = Path(file.filename or "").suffix.lower()
        if suffix not in self.supported_suffixes:
            raise UnsupportedDocumentError(f"Unsupported file type: {suffix or 'unknown'}")

        content = await file.read()
        if suffix in {".txt", ".md"}:
            return self._decode_text(content)
        if suffix == ".csv":
            return self._extract_csv(content)
        if suffix == ".json":
            return self._extract_json(content)
        if suffix == ".pdf":
            return self._extract_pdf(content)
        if suffix == ".docx":
            return self._extract_docx(content)
        if suffix == ".xlsx":
            return self._extract_xlsx(content)
        raise UnsupportedDocumentError(f"Unsupported file type: {suffix}")

    def _decode_text(self, content: bytes) -> str:
        for encoding in ("utf-8", "utf-16", "cp1252"):
            try:
                return content.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="ignore").strip()

    def _extract_csv(self, content: bytes) -> str:
        decoded = self._decode_text(content)
        rows = csv.reader(io.StringIO(decoded))
        return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows)

    def _extract_json(self, content: bytes) -> str:
        decoded = self._decode_text(content)
        data = json.loads(decoded)
        return json.dumps(data, indent=2, sort_keys=True)

    def _extract_pdf(self, content: bytes) -> str:
        try:
            import pdfplumber
        except ImportError as exc:
            raise UnsupportedDocumentError("Install the 'documents' extra to parse PDF files") from exc

        try:
            text_parts: list[str] = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    text_parts.append(page.extract_text() or "")
            return "\n".join(text_parts).strip()
        except Exception as exc:
            raise UnsupportedDocumentError("Unable to parse PDF document") from exc

    def _extract_docx(self, content: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise UnsupportedDocumentError("Install the 'documents' extra to parse DOCX files") from exc

        try:
            document = Document(io.BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
            table_rows: list[str] = []
            for table in document.tables:
                for row in table.rows:
                    table_rows.append(" | ".join(cell.text.strip() for cell in row.cells))
            return "\n".join(paragraphs + table_rows).strip()
        except Exception as exc:
            raise UnsupportedDocumentError("Unable to parse DOCX document") from exc

    def _extract_xlsx(self, content: bytes) -> str:
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise UnsupportedDocumentError("Install the 'documents' extra to parse XLSX files") from exc

        try:
            workbook = load_workbook(io.BytesIO(content), data_only=True)
            rows: list[str] = []
            for sheet in workbook.worksheets:
                rows.append(f"Sheet: {sheet.title}")
                for values in sheet.iter_rows(values_only=True):
                    cells = ["" if value is None else str(value).strip() for value in values]
                    if any(cells):
                        rows.append(" | ".join(cells))
            return "\n".join(rows).strip()
        except Exception as exc:
            raise UnsupportedDocumentError("Unable to parse XLSX document") from exc
